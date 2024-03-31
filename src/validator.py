import logging

import jinja2
from docx import Document

from src.base import AbstractValidator


class DocumentValidator(AbstractValidator):
    def __init__(self, template_path, dataset):
        self.template_path = template_path
        self.dataset = dataset
        self.env = jinja2.Environment()
        self.template_content = self.extract_text()

    def extract_text(self):
        doc = Document(self.template_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)

    def process(self):
        self.check_syntax()
        return self.validate()

    def check_syntax(self):
        try:
            self.env.parse(self.template_content)
            logging.info('No syntax errors')
        except jinja2.TemplateSyntaxError as e:
            raise SyntaxError

    def validate(self):
        ast = self.env.parse(self.template_content)
        variables = {node for node in jinja2.meta.find_undeclared_variables(ast)}
        logging.info(f'Template {str(self.template_path)}, variables {str(variables)}')
        difference = variables - set(self.dataset.keys())
        if difference:
            logging.error(f'Not enough data, {str(difference)}')
            raise AttributeError
        return variables
